import csv
import random
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime as dt

WELCOME_TEXT = (
    "We bid thee a most splendid welcome to the nuptials of Nicki & Tony. "
    "The ceremony hath concluded, and now the grand revelry begins! "
    "Shouldst thou yearn for any manner of tipple, pray complete the tasks "
    "inscribed upon this card and deliver it into the capable hands of one of "
    "the gracious sisters (Claire H or Rachel T). In exchange, thou shalt "
    "receive a well-earned token, granting thee a free drink to gladden thy heart.\n\n"
    "But, above all else, rejoice and make merry! We are overjoyed by thy presence; "
    "feast as thou wilt, and partake of wine to thy contentment. May laughter fill the halls, "
    "and may this celebration bring thee everlasting delight!\n\n"
    "Use the QR code on the back to upload your images"
)

def get_data_from_csv(file_path):
    """
    Reads a local CSV file and returns a list of dictionaries with:
      {
        "question": <str>,
        "remaining_uses": <int>
      }
    Assumes column headers:
      - "Question"
      - "Number of Uses"
    """
    data = []
    with open(file_path, mode="r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            q_text = row.get("Question", "").strip()
            usage_str = row.get("Number of Uses", "0").strip()
            try:
                usage_limit = int(usage_str)
            except ValueError:
                usage_limit = 0

            if q_text and usage_limit > 0:
                data.append({
                    "question": q_text,
                    "remaining_uses": usage_limit
                })
    return data

def pick_question(question_pool, used_texts):
    """
    Randomly pick one question that:
      - still has remaining_uses > 0,
      - not already in used_texts (avoid duplicates on the same page).
    Decrements usage by 1.
    Returns question text or None if none available.
    """
    valid = [
        q for q in question_pool
        if q["remaining_uses"] > 0 and q["question"] not in used_texts
    ]
    if not valid:
        return None

    chosen = random.choice(valid)
    chosen["remaining_uses"] -= 1
    return chosen["question"]

def main():
    # CSV sources
    CSV_FILE_1 = r"./jukebox2/useful tools/I spy Data - Picture.csv"   # Pool 1
    CSV_FILE_2 = r"./jukebox2/useful tools/I spy Data - Questions.csv" # Pool 2

    question_pool_1 = get_data_from_csv(CSV_FILE_1)
    question_pool_2 = get_data_from_csv(CSV_FILE_2)

    # Randomize question order (optional)
    random.shuffle(question_pool_1)
    random.shuffle(question_pool_2)

    doc = Document()

    # Configure A5 landscape, zero margins
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(148)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(2)
    section.right_margin = Mm(0)

    # Remove any existing header/footer text
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    num_variations = 68
    questions_from_pool_1 = 6
    questions_from_pool_2 = 1

    for _ in range(num_variations):
        # --- Add a heading above the table ---
        heading_para = doc.add_paragraph("\n\n\tThy Quest For Further Tipples Begins!\t\t\tYour Quest Instructions!")
        heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_run = heading_para.runs[0]
        heading_run.font.name = "Segoe Script"
        heading_run.font.size = Pt(11)
        heading_run.bold = True

        # Prepare list of 7 questions (6 from pool_1, 1 from pool_2)
        selected = []
        used_texts = set()

        # Pick 6 from pool_1
        for __ in range(questions_from_pool_1):
            q = pick_question(question_pool_1, used_texts)
            if not q:
                selected.append("Not enough unique questions remain in Pool 1!")
                break
            selected.append(q)
            used_texts.add(q)

        # Pick 1 from pool_2
        q2 = pick_question(question_pool_2, used_texts)
        if not q2:
            selected.append("Not enough unique questions remain in Pool 2!")
        else:
            selected.append(q2)
            used_texts.add(q2)

        # Ensure total 7 items
        while len(selected) < 7:
            selected.append("")

        # Create a 7-row x 2-column table
        table = doc.add_table(rows=7, cols=2)
        
        # 1) Make the table span the entire width of the page
        #    Start at left edge, fill the whole 210mm width
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        table.allow_autofit = False
        table.autofit = False
        
        # Force each column to be 105 mm => total 210 mm
        table.columns[0].width = Mm(103)
        table.columns[1].width = Mm(105)

        # Merge left cells (rows 0..6) into one large cell
        # to contain WELCOME_TEXT
        top_left_cell = table.cell(0, 0)
        bottom_left_cell = table.cell(6, 0)
        merged_left = top_left_cell.merge(bottom_left_cell)

        # Put the WELCOME_TEXT in that merged left cell
        p_left = merged_left.paragraphs[0]
        p_left.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_left = p_left.add_run(WELCOME_TEXT)
        run_left.font.name = "Alex Brush"
        run_left.font.size = Pt(11)

        # Fill the 7 questions on the right column (rows 0..6)
        for row_idx in range(7):
            cell_q = table.cell(row_idx, 1)
            paragraph_q = cell_q.paragraphs[0]
            paragraph_q.style = doc.styles['List Bullet']
            paragraph_q.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run_q = paragraph_q.add_run(selected[row_idx])
            run_q.font.name = "Alex Brush"
            run_q.font.size = Pt(10)

        # End this variation with a page break
        doc.add_page_break()

    doc_name = f"A5_Landscape_{dt.now().strftime('%Y%m%d%H%M')}.docx"
    doc.save(doc_name)
    print(f"Document generated: {doc_name}")

if __name__ == "__main__":
    main()
